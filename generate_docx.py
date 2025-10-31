#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def create_xiyouji_docx():
    """创建包含《西游记》介绍的Word文档"""
    
    # 确保目标目录存在
    output_path = "/Users/issuser/Documents/file-test/sample.docx"
    output_dir = os.path.dirname(output_path)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # 创建文档
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('《西游记》介绍', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 基本信息
    heading1 = doc.add_heading('一、基本信息', 1)
    
    info_text = """《西游记》是中国古代第一部浪漫主义章回体长篇神魔小说，是中国古典四大名著之一。全书共一百回，主要描写了孙悟空、猪八戒、沙僧三人保护唐僧西行取经，沿途遇到八十一难，一路降妖伏魔，化险为夷，最后到达西天、取得真经的故事。"""
    
    p1 = doc.add_paragraph(info_text)
    p1.paragraph_format.first_line_indent = Inches(0.5)
    p1.paragraph_format.line_spacing = 1.5
    
    # 作者简介
    heading2 = doc.add_heading('二、作者简介', 1)
    
    author_text = """作者：吴承恩（约1500年-1582年），字汝忠，号射阳山人，淮安府山阳县（今江苏省淮安市淮安区）人，明代著名小说家。吴承恩自幼聪慧，博览群书，尤其喜爱神话故事。他科举不顺，仕途坎坷，晚年才创作出这部伟大的神魔小说。"""
    
    p2 = doc.add_paragraph(author_text)
    p2.paragraph_format.first_line_indent = Inches(0.5)
    p2.paragraph_format.line_spacing = 1.5
    
    # 主要人物
    heading3 = doc.add_heading('三、主要人物', 1)
    
    characters = [
        ("孙悟空", "又称齐天大圣、斗战胜佛，是花果山灵石孕育而生的石猴。他桀骜不驯，神通广大，会七十二般变化，有一双火眼金睛，手持如意金箍棒。曾大闹天宫，后被如来佛祖压在五行山下五百年。后经观音菩萨点化，保护唐僧西天取经，一路斩妖除魔，立下赫赫战功，最终修成正果，被封为斗战胜佛。"),
        ("唐僧", "俗家姓陈，法名玄奘，法号三藏，被尊称为唐僧。他是如来佛祖的二徒弟金蝉子转世，一心向佛，慈悲为怀。奉唐太宗之命西天取经，历经九九八十一难，最终取得真经，被封为旃檀功德佛。"),
        ("猪八戒", "原为天蓬元帅，因调戏嫦娥被贬下凡，错投猪胎。法号悟能，是唐僧的二徒弟。他好吃懒做，贪财好色，但憨厚忠诚，在取经路上也立下不少功劳。最终被封为净坛使者。"),
        ("沙僧", "原为天庭卷帘大将，因打破琉璃盏被贬流沙河。法号悟净，是唐僧的三徒弟。他任劳任怨，忠心耿耿，虽然本领不高，但勤勤恳恳保护师父。最终被封为金身罗汉。"),
    ]
    
    for name, desc in characters:
        p = doc.add_paragraph()
        run_name = p.add_run(name + "：")
        run_name.bold = True
        run_name.font.size = Pt(12)
        run_desc = p.add_run(desc)
        run_desc.font.size = Pt(11)
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.line_spacing = 1.5
    
    # 主要内容
    heading4 = doc.add_heading('四、主要内容', 1)
    
    content_parts = [
        ("第一部分（1-7回）", "孙悟空的出世、拜师学艺、大闹天宫，最后被如来佛祖压在五行山下。"),
        ("第二部分（8-12回）", "交代取经的缘由，唐僧的身世，以及观音菩萨点化孙悟空、猪八戒、沙僧等情节。"),
        ("第三部分（13-100回）", "唐僧师徒四人西天取经的主要过程，历经九九八十一难，降伏各路妖魔，最终取得真经，修成正果。"),
    ]
    
    for title, desc in content_parts:
        p = doc.add_paragraph()
        run_title = p.add_run(title + "：")
        run_title.bold = True
        run_desc = p.add_run(desc)
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.line_spacing = 1.5
    
    # 艺术特色
    heading5 = doc.add_heading('五、艺术特色', 1)
    
    features = [
        "想象奇特，情节生动，充满浪漫主义色彩",
        "人物形象鲜明，性格各异，栩栩如生",
        "语言生动活泼，幽默风趣，雅俗共赏",
        "寓意深刻，既有神话的奇幻，又有现实的讽刺",
        "结构宏大，故事连贯，主次分明",
    ]
    
    for feature in features:
        p = doc.add_paragraph(feature, style='List Bullet')
        p.paragraph_format.line_spacing = 1.5
    
    # 影响与地位
    heading6 = doc.add_heading('六、影响与地位', 1)
    
    influence_text = """《西游记》自问世以来，在中国乃至世界都产生了深远的影响。它不仅是中国古典文学的巅峰之作，也是世界文学宝库中的瑰宝。该书被翻译成多种语言，在世界范围内广为流传。其故事情节和人物形象深入人心，被多次改编成戏曲、电影、电视剧、动画片等各种艺术形式，成为中华文化的重要符号之一。

书中蕴含的不畏艰难、坚持不懈、团结协作的精神，以及对善恶是非的思考，对后世产生了深刻的教育意义和启迪作用。《西游记》不仅是一部文学经典，更是一座精神丰碑，代表着中华民族的智慧和创造力。"""
    
    p_influence = doc.add_paragraph(influence_text)
    p_influence.paragraph_format.first_line_indent = Inches(0.5)
    p_influence.paragraph_format.line_spacing = 1.5
    
    # 保存文档
    doc.save(output_path)
    print(f"Word文档已成功创建：{output_path}")
    print("文档内容：《西游记》介绍")

if __name__ == "__main__":
    create_xiyouji_docx()
